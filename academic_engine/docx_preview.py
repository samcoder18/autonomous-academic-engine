"""Shared DOCX preview formatter for work-specific thesis bundles."""

from __future__ import annotations

import argparse
import sys
import tomllib
from collections.abc import Sequence
from dataclasses import dataclass, field
from pathlib import Path
from types import SimpleNamespace
from typing import Any, TextIO

from .workspace import WorkspaceConfigError, load_work_config, load_workspace_config

FONT_NAME = "Times New Roman"


class DocxPreviewError(RuntimeError):
    """Raised when DOCX preview formatting cannot proceed."""


class DocxPreviewDependencyError(DocxPreviewError):
    """Raised when optional DOCX formatting dependencies are missing."""


@dataclass(frozen=True)
class DocxPreviewConfig:
    title: str
    subject: str
    author: str
    keywords: str
    major_titles: tuple[str, ...]
    title_center_until: int
    title_right_align_indices: tuple[int, ...]
    title_spacing: dict[int, tuple[int, int]] = field(default_factory=dict)
    format_contents_table: bool = False
    compact_style_justified: bool = False


def load_docx_preview_config(root_dir: str | Path, work_id: str | None = None) -> DocxPreviewConfig:
    workspace = load_workspace_config(root_dir)
    work = load_work_config(workspace, work_id or workspace.default_work)
    if not work.thesis:
        raise WorkspaceConfigError(f"Work `{work.slug}` не поддерживает thesis lane.")

    work_file = work.work_dir / "work.toml"
    with work_file.open("rb") as handle:
        payload = tomllib.load(handle)

    thesis = payload.get("thesis")
    if not isinstance(thesis, dict):
        raise WorkspaceConfigError(f"Секция [thesis] в {work_file} должна быть таблицей.")
    raw = thesis.get("docx_preview")
    if not isinstance(raw, dict):
        raise WorkspaceConfigError(f"В {work_file} отсутствует секция [thesis.docx_preview].")
    if raw.get("enabled") is False:
        raise WorkspaceConfigError(f"DOCX preview disabled for work `{work.slug}`.")

    return DocxPreviewConfig(
        title=work.title,
        subject=_required_text(raw, "subject", work_file),
        author=_required_text(raw, "author", work_file),
        keywords=_required_text(raw, "keywords", work_file),
        major_titles=_required_text_tuple(raw, "major_titles", work_file),
        title_center_until=_required_int(raw, "title_center_until", work_file),
        title_right_align_indices=_required_int_tuple(raw, "title_right_align_indices", work_file),
        title_spacing=_title_spacing(raw.get("title_spacing"), work_file),
        format_contents_table=bool(raw.get("format_contents_table", False)),
        compact_style_justified=bool(raw.get("compact_style_justified", False)),
    )


def format_docx_preview(source: Path, target: Path, config: DocxPreviewConfig) -> Path:
    if not source.exists():
        raise DocxPreviewError(f"Input DOCX not found: {source}")

    deps = _load_docx_dependencies()
    body_size = deps.Pt(14)

    def set_run_font(run: Any, size: Any = body_size, *, bold: bool | None = None) -> None:
        run.font.name = FONT_NAME
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(deps.qn(f"w:{attr}"), FONT_NAME)
        run.font.size = size
        run.font.color.rgb = deps.RGBColor(0, 0, 0)
        if bold is not None:
            run.bold = bold

    def set_style_font(style: Any, size: Any, *, bold: bool | None = None) -> None:
        style.font.name = FONT_NAME
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(deps.qn(f"w:{attr}"), FONT_NAME)
        style.font.size = size
        style.font.color.rgb = deps.RGBColor(0, 0, 0)
        if bold is not None:
            style.font.bold = bold

    def configure_defaults(document: Any) -> None:
        styles = document.styles.element
        defaults = styles.find(deps.qn("w:docDefaults"))
        if defaults is None:
            defaults = deps.OxmlElement("w:docDefaults")
            styles.insert(0, defaults)

        rpr_default = defaults.find(deps.qn("w:rPrDefault"))
        if rpr_default is None:
            rpr_default = deps.OxmlElement("w:rPrDefault")
            defaults.append(rpr_default)
        rpr = rpr_default.find(deps.qn("w:rPr"))
        if rpr is None:
            rpr = deps.OxmlElement("w:rPr")
            rpr_default.append(rpr)
        fonts = rpr.find(deps.qn("w:rFonts"))
        if fonts is None:
            fonts = deps.OxmlElement("w:rFonts")
            rpr.append(fonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(deps.qn(f"w:{attr}"), FONT_NAME)
        for tag in ("w:sz", "w:szCs"):
            size = rpr.find(deps.qn(tag))
            if size is None:
                size = deps.OxmlElement(tag)
                rpr.append(size)
            size.set(deps.qn("w:val"), "28")

        ppr_default = defaults.find(deps.qn("w:pPrDefault"))
        if ppr_default is None:
            ppr_default = deps.OxmlElement("w:pPrDefault")
            defaults.append(ppr_default)
        ppr = ppr_default.find(deps.qn("w:pPr"))
        if ppr is None:
            ppr = deps.OxmlElement("w:pPr")
            ppr_default.append(ppr)
        spacing = ppr.find(deps.qn("w:spacing"))
        if spacing is None:
            spacing = deps.OxmlElement("w:spacing")
            ppr.append(spacing)
        spacing.set(deps.qn("w:line"), "360")
        spacing.set(deps.qn("w:lineRule"), "auto")
        spacing.set(deps.qn("w:before"), "0")
        spacing.set(deps.qn("w:after"), "0")

    def configure_styles(document: Any) -> None:
        for name in ("Normal", "Body Text", "First Paragraph"):
            if name not in document.styles:
                continue
            style = document.styles[name]
            set_style_font(style, body_size)
            fmt = style.paragraph_format
            fmt.alignment = deps.WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.first_line_indent = deps.Cm(1.25)
            fmt.line_spacing = 1.5
            fmt.space_before = deps.Pt(0)
            fmt.space_after = deps.Pt(0)

        if "Compact" in document.styles:
            style = document.styles["Compact"]
            set_style_font(style, body_size)
            fmt = style.paragraph_format
            if config.compact_style_justified:
                fmt.alignment = deps.WD_ALIGN_PARAGRAPH.JUSTIFY
                fmt.first_line_indent = deps.Cm(1.25)
            fmt.line_spacing = 1.5
            fmt.space_before = deps.Pt(0)
            fmt.space_after = deps.Pt(0)

        for name, alignment, before, after in (
            ("Heading 1", deps.WD_ALIGN_PARAGRAPH.CENTER, 0, 12),
            ("Heading 2", deps.WD_ALIGN_PARAGRAPH.LEFT, 12, 6),
        ):
            if name not in document.styles:
                continue
            style = document.styles[name]
            set_style_font(style, body_size, bold=True)
            fmt = style.paragraph_format
            fmt.alignment = alignment
            fmt.first_line_indent = deps.Cm(0)
            fmt.line_spacing = 1.5
            fmt.space_before = deps.Pt(before)
            fmt.space_after = deps.Pt(after)
            fmt.keep_with_next = True

        if "Footnote Text" in document.styles:
            style = document.styles["Footnote Text"]
            set_style_font(style, deps.Pt(10))
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.space_before = deps.Pt(0)
            style.paragraph_format.space_after = deps.Pt(0)
        if "Footnote Reference" in document.styles:
            set_style_font(document.styles["Footnote Reference"], deps.Pt(10))

    def add_page_field(paragraph: Any) -> None:
        paragraph.alignment = deps.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        set_run_font(run, deps.Pt(12))
        begin = deps.OxmlElement("w:fldChar")
        begin.set(deps.qn("w:fldCharType"), "begin")
        instruction = deps.OxmlElement("w:instrText")
        instruction.set(deps.qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        separate = deps.OxmlElement("w:fldChar")
        separate.set(deps.qn("w:fldCharType"), "separate")
        value = deps.OxmlElement("w:t")
        value.text = "2"
        end = deps.OxmlElement("w:fldChar")
        end.set(deps.qn("w:fldCharType"), "end")
        run._r.extend((begin, instruction, separate, value, end))

    def configure_sections(document: Any) -> None:
        for section in document.sections:
            section.page_width = deps.Mm(210)
            section.page_height = deps.Mm(297)
            section.left_margin = deps.Mm(30)
            section.right_margin = deps.Mm(20)
            section.top_margin = deps.Mm(20)
            section.bottom_margin = deps.Mm(20)
            section.header_distance = deps.Mm(12.5)
            section.footer_distance = deps.Mm(12.5)
            section.different_first_page_header_footer = True
            section.footer.is_linked_to_previous = False
            paragraph = section.footer.paragraphs[0]
            paragraph.clear()
            add_page_field(paragraph)

    def configure_numbering(document: Any) -> None:
        numbering = document.part.numbering_part.element
        for level in numbering.findall(".//w:lvl[@w:ilvl='0']", numbering.nsmap):
            number_format = level.find("w:numFmt", numbering.nsmap)
            if number_format is None or number_format.get(deps.qn("w:val")) != "decimal":
                continue
            suffix = level.find("w:suff", numbering.nsmap)
            if suffix is None:
                suffix = deps.OxmlElement("w:suff")
                level.insert(3, suffix)
            suffix.set(deps.qn("w:val"), "space")
            ppr = level.find("w:pPr", numbering.nsmap)
            if ppr is None:
                ppr = deps.OxmlElement("w:pPr")
                level.append(ppr)
            indentation = ppr.find("w:ind", numbering.nsmap)
            if indentation is None:
                indentation = deps.OxmlElement("w:ind")
                ppr.append(indentation)
            indentation.set(deps.qn("w:left"), "540")
            indentation.set(deps.qn("w:hanging"), "360")

    def configure_contents_table(document: Any) -> None:
        if not document.tables:
            return
        table = document.tables[0]
        for row in table.rows:
            row.height = None
            for cell in row.cells:
                margins = cell._tc.get_or_add_tcPr().find(deps.qn("w:tcMar"))
                if margins is None:
                    margins = deps.OxmlElement("w:tcMar")
                    cell._tc.get_or_add_tcPr().append(margins)
                for side, value in (("top", "25"), ("bottom", "25"), ("left", "50"), ("right", "50")):
                    node = margins.find(deps.qn(f"w:{side}"))
                    if node is None:
                        node = deps.OxmlElement(f"w:{side}")
                        margins.append(node)
                    node.set(deps.qn("w:w"), value)
                    node.set(deps.qn("w:type"), "dxa")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = deps.Cm(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = deps.Pt(0)
                    paragraph.paragraph_format.space_after = deps.Pt(0)
                    for run in paragraph.runs:
                        set_run_font(run, deps.Pt(11))

    def configure_paragraphs(document: Any) -> None:
        major_titles = set(config.major_titles)
        right_align = set(config.title_right_align_indices)

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            for run in paragraph.runs:
                set_run_font(run, body_size, bold=paragraph.style.name in ("Heading 1", "Heading 2"))
            if text in major_titles:
                paragraph.paragraph_format.page_break_before = True
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.alignment = deps.WD_ALIGN_PARAGRAPH.CENTER
            if 0 <= index < config.title_center_until:
                paragraph.paragraph_format.first_line_indent = deps.Cm(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = deps.Pt(0)
                paragraph.paragraph_format.space_after = deps.Pt(0)
                paragraph.alignment = deps.WD_ALIGN_PARAGRAPH.CENTER

        for index, (before, after) in config.title_spacing.items():
            if index >= len(document.paragraphs):
                continue
            paragraph = document.paragraphs[index]
            paragraph.paragraph_format.space_before = deps.Pt(before)
            paragraph.paragraph_format.space_after = deps.Pt(after)

        for index in right_align:
            if index < len(document.paragraphs):
                document.paragraphs[index].alignment = deps.WD_ALIGN_PARAGRAPH.RIGHT

    document = deps.Document(source)
    configure_defaults(document)
    configure_styles(document)
    configure_sections(document)
    configure_numbering(document)
    if config.format_contents_table:
        configure_contents_table(document)
    configure_paragraphs(document)

    document.core_properties.title = config.title
    document.core_properties.subject = config.subject
    document.core_properties.author = config.author
    document.core_properties.keywords = config.keywords

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return target


def main(
    argv: Sequence[str] | None = None,
    *,
    root_dir: str | Path | None = None,
    stdout: TextIO | None = None,
    stderr: TextIO | None = None,
) -> int:
    out = stdout or sys.stdout
    err = stderr or sys.stderr
    parser = argparse.ArgumentParser(description="Render a formatted DOCX preview for a thesis work.")
    parser.add_argument("--work", help="Workspace work slug. Defaults to workspace.toml default_work.")
    parser.add_argument("--input", dest="input_docx", help="Input DOCX path. Defaults to the work export path.")
    parser.add_argument("--output", dest="output_docx", help="Output DOCX path. Defaults to *-preview.docx.")
    args = parser.parse_args(argv)

    root = Path(root_dir or ".").expanduser().resolve()
    try:
        workspace = load_workspace_config(root)
        work = load_work_config(workspace, args.work or workspace.default_work)
        if not work.thesis:
            raise WorkspaceConfigError(f"Work `{work.slug}` не поддерживает thesis lane.")
        config = load_docx_preview_config(root, work.slug)
        source = _resolve_cli_path(root, args.input_docx) if args.input_docx else work.thesis.export_docx_path
        if args.output_docx:
            target = _resolve_cli_path(root, args.output_docx)
        else:
            target = work.thesis.paths.output_docx_dir / f"{source.stem}-preview.docx"
        result = format_docx_preview(source, target, config)
    except (DocxPreviewError, WorkspaceConfigError) as exc:
        print(str(exc), file=err)
        return 1

    print(result, file=out)
    return 0


def _load_docx_dependencies() -> SimpleNamespace:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Mm, Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise DocxPreviewDependencyError(
            "python-docx is required for DOCX preview formatting. Install the optional `python-docx` package."
        ) from exc

    return SimpleNamespace(
        Document=Document,
        WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
        OxmlElement=OxmlElement,
        qn=qn,
        Cm=Cm,
        Mm=Mm,
        Pt=Pt,
        RGBColor=RGBColor,
    )


def _resolve_cli_path(root: Path, raw: str) -> Path:
    path = Path(raw).expanduser()
    if path.is_absolute():
        return path
    return root / path


def _required_text(payload: dict[str, Any], key: str, source: Path) -> str:
    value = payload.get(key)
    if not isinstance(value, str) or not value.strip():
        raise WorkspaceConfigError(f"В {source} docx_preview.{key} должен быть непустой строкой.")
    return value.strip()


def _required_int(payload: dict[str, Any], key: str, source: Path) -> int:
    value = payload.get(key)
    if not isinstance(value, int):
        raise WorkspaceConfigError(f"В {source} docx_preview.{key} должен быть целым числом.")
    return value


def _required_text_tuple(payload: dict[str, Any], key: str, source: Path) -> tuple[str, ...]:
    value = payload.get(key)
    if not isinstance(value, list) or not value or not all(isinstance(item, str) and item.strip() for item in value):
        raise WorkspaceConfigError(f"В {source} docx_preview.{key} должен быть непустым списком строк.")
    return tuple(item.strip() for item in value)


def _required_int_tuple(payload: dict[str, Any], key: str, source: Path) -> tuple[int, ...]:
    value = payload.get(key)
    if not isinstance(value, list) or not all(isinstance(item, int) for item in value):
        raise WorkspaceConfigError(f"В {source} docx_preview.{key} должен быть списком целых чисел.")
    return tuple(value)


def _title_spacing(raw: object, source: Path) -> dict[int, tuple[int, int]]:
    if not isinstance(raw, list):
        raise WorkspaceConfigError(f"В {source} docx_preview.title_spacing должен быть списком таблиц.")
    result: dict[int, tuple[int, int]] = {}
    for item in raw:
        if not isinstance(item, dict):
            raise WorkspaceConfigError(f"В {source} каждый элемент docx_preview.title_spacing должен быть таблицей.")
        index = _required_int(item, "index", source)
        before = _required_int(item, "before_pt", source)
        after = _required_int(item, "after_pt", source)
        result[index] = (before, after)
    return result


if __name__ == "__main__":
    raise SystemExit(main())
