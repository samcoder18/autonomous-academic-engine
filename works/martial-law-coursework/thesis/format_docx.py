from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(14)


def set_run_font(run, size: Pt = BODY_SIZE, *, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT_NAME)
    run.font.size = size
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def set_style_font(style, size: Pt, *, bold: bool | None = None) -> None:
    style.font.name = FONT_NAME
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT_NAME)
    style.font.size = size
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold


def configure_defaults(document: Document) -> None:
    styles = document.styles.element
    defaults = styles.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles.insert(0, defaults)

    rpr_default = defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT_NAME)
    for tag in ("w:sz", "w:szCs"):
        size = rpr.find(qn(tag))
        if size is None:
            size = OxmlElement(tag)
            rpr.append(size)
        size.set(qn("w:val"), "28")

    ppr_default = defaults.find(qn("w:pPrDefault"))
    if ppr_default is None:
        ppr_default = OxmlElement("w:pPrDefault")
        defaults.append(ppr_default)
    ppr = ppr_default.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        ppr_default.append(ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")


def configure_styles(document: Document) -> None:
    for name in ("Normal", "Body Text", "First Paragraph", "Compact"):
        if name not in document.styles:
            continue
        style = document.styles[name]
        set_style_font(style, BODY_SIZE)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Cm(1.25)
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)

    for name, alignment, before, after in (
        ("Heading 1", WD_ALIGN_PARAGRAPH.CENTER, 0, 12),
        ("Heading 2", WD_ALIGN_PARAGRAPH.LEFT, 12, 6),
    ):
        if name not in document.styles:
            continue
        style = document.styles[name]
        set_style_font(style, BODY_SIZE, bold=True)
        fmt = style.paragraph_format
        fmt.alignment = alignment
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.keep_with_next = True

    if "Footnote Text" in document.styles:
        style = document.styles["Footnote Text"]
        set_style_font(style, Pt(10))
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
    if "Footnote Reference" in document.styles:
        set_style_font(document.styles["Footnote Reference"], Pt(10))


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, Pt(12))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "2"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, value, end))


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.header_distance = Mm(12.5)
        section.footer_distance = Mm(12.5)
        section.different_first_page_header_footer = True
        section.footer.is_linked_to_previous = False
        paragraph = section.footer.paragraphs[0]
        paragraph.clear()
        add_page_field(paragraph)


def configure_numbering(document: Document) -> None:
    numbering = document.part.numbering_part.element
    for level in numbering.findall(".//w:lvl[@w:ilvl='0']", numbering.nsmap):
        number_format = level.find("w:numFmt", numbering.nsmap)
        if number_format is None or number_format.get(qn("w:val")) != "decimal":
            continue
        suffix = level.find("w:suff", numbering.nsmap)
        if suffix is None:
            suffix = OxmlElement("w:suff")
            level.insert(3, suffix)
        suffix.set(qn("w:val"), "space")
        ppr = level.find("w:pPr", numbering.nsmap)
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            level.append(ppr)
        indentation = ppr.find("w:ind", numbering.nsmap)
        if indentation is None:
            indentation = OxmlElement("w:ind")
            ppr.append(indentation)
        indentation.set(qn("w:left"), "540")
        indentation.set(qn("w:hanging"), "360")


def configure_contents_table(document: Document) -> None:
    if not document.tables:
        return
    table = document.tables[0]
    for row in table.rows:
        row.height = None
        for cell in row.cells:
            margins = cell._tc.get_or_add_tcPr().find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(margins)
            for side, value in (("top", "25"), ("bottom", "25"), ("left", "50"), ("right", "50")):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, Pt(11))


def configure_paragraphs(document: Document) -> None:
    section_titles = {
        "Содержание",
        "Введение",
        "Глава 1. Военное положение как особый государственно-правовой режим",
        "Глава 2. Конституционно-правовые основания и процедура введения военного положения",
        "Глава 3. Реализация режима военного положения и гарантии законности",
        "Заключение",
        "Список использованных источников",
    }

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for run in paragraph.runs:
            if paragraph.style.name in ("Heading 1", "Heading 2"):
                set_run_font(run, BODY_SIZE, bold=True)
            else:
                set_run_font(run, BODY_SIZE)
        if text in section_titles:
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # The title page is centered except for the author and reviewer blocks.
    for index, paragraph in enumerate(document.paragraphs[:14]):
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if index in (8, 9, 10, 11):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title_spacing = {
        0: (0, 10),
        2: (0, 70),
        3: (0, 8),
        6: (8, 75),
        8: (0, 18),
        10: (0, 18),
        12: (0, 105),
    }
    for index, (before, after) in title_spacing.items():
        if index < len(document.paragraphs):
            paragraph = document.paragraphs[index]
            paragraph.paragraph_format.space_before = Pt(before)
            paragraph.paragraph_format.space_after = Pt(after)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: format_docx.py INPUT.docx OUTPUT.docx", file=sys.stderr)
        return 2

    source = Path(sys.argv[1]).resolve()
    target = Path(sys.argv[2]).resolve()
    document = Document(source)
    configure_defaults(document)
    configure_styles(document)
    configure_sections(document)
    configure_numbering(document)
    configure_contents_table(document)
    configure_paragraphs(document)

    document.core_properties.title = (
        "Военное положение: конституционно-правовые основы и процедура введения"
    )
    document.core_properties.subject = "Курсовая работа по теории государства и права"
    document.core_properties.author = "ФИО СТУДЕНТА"
    document.core_properties.keywords = (
        "военное положение; Конституция Российской Федерации; особый правовой режим"
    )

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
