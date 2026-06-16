from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(14)


def set_run_font(run, size: Pt = BODY_SIZE, *, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def set_style_font(style, size: Pt, *, bold: bool | None = None) -> None:
    style.font.name = FONT_NAME
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_NAME)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_NAME)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = size
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold


def set_doc_defaults(document: Document) -> None:
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)

    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT_NAME)

    for tag in ("w:sz", "w:szCs"):
        size = rpr.find(qn(tag))
        if size is None:
            size = OxmlElement(tag)
            rpr.append(size)
        size.set(qn("w:val"), "28")

    ppr_default = doc_defaults.find(qn("w:pPrDefault"))
    if ppr_default is None:
        ppr_default = OxmlElement("w:pPrDefault")
        doc_defaults.append(ppr_default)
    ppr = ppr_default.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        ppr_default.append(ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")


def configure_styles(document: Document) -> None:
    for style_name in ("Normal", "Body Text", "First Paragraph"):
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        set_style_font(style, BODY_SIZE)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Cm(1.25)
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)

    if "Compact" in document.styles:
        style = document.styles["Compact"]
        set_style_font(style, BODY_SIZE)
        fmt = style.paragraph_format
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)

    if "Heading 1" in document.styles:
        style = document.styles["Heading 1"]
        set_style_font(style, BODY_SIZE, bold=True)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(12)
        fmt.keep_with_next = True

    if "Heading 2" in document.styles:
        style = document.styles["Heading 2"]
        set_style_font(style, BODY_SIZE, bold=True)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)
        fmt.keep_with_next = True

    for style_name in ("Footnote Text", "Footnote Reference"):
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        set_style_font(style, Pt(10))
        if style_name == "Footnote Text":
            fmt = style.paragraph_format
            fmt.line_spacing = 1.0
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, Pt(12))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "2"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, value, end))


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.header_distance = Mm(12.5)
        section.footer_distance = Mm(12.5)
        section.different_first_page_header_footer = True

        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        add_page_field(paragraph)


def configure_numbering(document: Document) -> None:
    numbering = document.part.numbering_part.element
    for level in numbering.findall(".//w:lvl[@w:ilvl='0']", numbering.nsmap):
        number_format = level.find("w:numFmt", numbering.nsmap)
        if number_format is None or number_format.get(qn("w:val")) != "decimal":
            continue
        paragraph_properties = level.find("w:pPr", numbering.nsmap)
        if paragraph_properties is None:
            paragraph_properties = OxmlElement("w:pPr")
            level.append(paragraph_properties)
        suffix = level.find("w:suff", numbering.nsmap)
        if suffix is None:
            suffix = OxmlElement("w:suff")
            level.insert(3, suffix)
        suffix.set(qn("w:val"), "space")
        indentation = paragraph_properties.find("w:ind", numbering.nsmap)
        if indentation is None:
            indentation = OxmlElement("w:ind")
            paragraph_properties.append(indentation)
        indentation.set(qn("w:left"), "540")
        indentation.set(qn("w:hanging"), "360")


def configure_paragraphs(document: Document) -> None:
    major_titles = {
        "Введение",
        "Глава 1. Конституционные поправки как теоретико-правовое явление",
        "Глава 2. Основные направления реализации поправок 2020 года в Российской Федерации",
        "Глава 3. Проблемы реализации конституционных поправок",
        "Заключение",
        "Список использованных источников",
    }

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        for run in paragraph.runs:
            if paragraph.style.name == "Heading 1":
                set_run_font(run, BODY_SIZE, bold=True)
            elif paragraph.style.name == "Heading 2":
                set_run_font(run, BODY_SIZE, bold=True)
            else:
                set_run_font(run, BODY_SIZE)

        if text in major_titles:
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if 0 <= index <= 11:
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_spacing = {
        0: (0, 10),
        2: (0, 70),
        3: (0, 8),
        6: (8, 75),
        7: (0, 0),
        8: (0, 18),
        9: (0, 0),
        10: (0, 105),
        11: (0, 0),
    }
    for index, (before, after) in title_spacing.items():
        if index >= len(document.paragraphs):
            continue
        paragraph = document.paragraphs[index]
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)

    for index in (7, 8, 9, 10):
        if index < len(document.paragraphs):
            document.paragraphs[index].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: format_docx.py INPUT.docx OUTPUT.docx", file=sys.stderr)
        return 2

    source = Path(sys.argv[1]).resolve()
    target = Path(sys.argv[2]).resolve()
    document = Document(source)

    set_doc_defaults(document)
    configure_styles(document)
    configure_sections(document)
    configure_numbering(document)
    configure_paragraphs(document)

    document.core_properties.title = "Проблемы реализации конституционных поправок"
    document.core_properties.subject = "Курсовая работа по теории государства и права"
    document.core_properties.author = "ФИО СТУДЕНТА"
    document.core_properties.keywords = "конституционные поправки; публичная власть; теория государства и права"

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
